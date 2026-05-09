from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from pathlib import Path

docx_path = 'collabdocs.docx'
doc = Document(docx_path)
output_dir = Path('docs/images')
output_dir.mkdir(parents=True, exist_ok=True)

image_count = 0

# Extract images from paragraphs
for para_idx, para in enumerate(doc.paragraphs):
    for run in para.runs:
        for rel in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            for inline in rel.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                embed_id = inline.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id:
                    image_part = run.part.related_parts[embed_id]
                    image_bytes = image_part.blob
                    # Determine file extension
                    ext = 'png'
                    if image_part.content_type == 'image/jpeg':
                        ext = 'jpg'
                    elif image_part.content_type == 'image/png':
                        ext = 'png'
                    
                    filename = f'image_{image_count}.{ext}'
                    filepath = output_dir / filename
                    with open(filepath, 'wb') as f:
                        f.write(image_bytes)
                    print(f'Extracted: {filename}')
                    image_count += 1

# Extract images from tables
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    for rel in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        for inline in rel.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                            embed_id = inline.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                image_part = run.part.related_parts[embed_id]
                                image_bytes = image_part.blob
                                ext = 'png'
                                if image_part.content_type == 'image/jpeg':
                                    ext = 'jpg'
                                
                                filename = f'image_{image_count}.{ext}'
                                filepath = output_dir / filename
                                with open(filepath, 'wb') as f:
                                    f.write(image_bytes)
                                print(f'Extracted: {filename}')
                                image_count += 1

print(f'\nTotal images extracted: {image_count}')
print(f'Saved to: {output_dir.absolute()}')
