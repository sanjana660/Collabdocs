from docx import Document
import os

docx_path = os.path.join(os.path.dirname(__file__), '..', 'collabdocs.docx')
doc = Document(docx_path)

print("=" * 80)
print("DOCX CONTENT")
print("=" * 80)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)

for table in doc.tables:
    print("\nTable:")
    for row in table.rows:
        print(" | ".join([cell.text for cell in row.cells]))
